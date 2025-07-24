# app/services/document_service.py
import os
import shutil
import aiofiles
from fastapi import HTTPException, UploadFile
from sqlalchemy.orm import Session
from sqlalchemy import or_, and_
from typing import Optional, List
from uuid import UUID
import magic
from datetime import datetime

from models import Document, User, DocumentPermission
from schemas import DocumentShare
from config import settings

class DocumentService:
    def __init__(self):
        self.upload_dir = settings.UPLOAD_DIR
        os.makedirs(self.upload_dir, exist_ok=True)

    async def upload_document(self, db: Session, file: UploadFile, user_id: UUID):
        """Upload và lưu trữ tài liệu"""
        
        # Validate file
        if not file.filename:
            raise HTTPException(status_code=400, detail="Tên file không hợp lệ")
        
        # Check file extension
        file_ext = file.filename.split('.')[-1].lower()
        if file_ext not in settings.ALLOWED_EXTENSIONS:
            raise HTTPException(
                status_code=400, 
                detail=f"Định dạng file không được hỗ trợ. Chỉ chấp nhận: {', '.join(settings.ALLOWED_EXTENSIONS)}"
            )
        
        # Check file size
        file.file.seek(0, 2)  # Move to end of file
        file_size = file.file.tell()
        file.file.seek(0)  # Reset to beginning
        
        if file_size > settings.MAX_FILE_SIZE:
            raise HTTPException(
                status_code=400, 
                detail=f"File quá lớn. Kích thước tối đa: {settings.MAX_FILE_SIZE / (1024*1024):.0f}MB"
            )
        
        # Create user directory
        user_dir = os.path.join(self.upload_dir, str(user_id))
        os.makedirs(user_dir, exist_ok=True)
        
        # Generate unique filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_filename = f"{timestamp}_{file.filename}"
        file_path = os.path.join(user_dir, safe_filename)
        
        try:
            # Save file
            async with aiofiles.open(file_path, 'wb') as f:
                content = await file.read()
                await f.write(content)
            
            # Detect MIME type
            mime_type = magic.from_file(file_path, mime=True)
            
            # Create document record
            document = Document(
                name=file.filename,
                original_name=file.filename,
                type=file_ext.upper(),
                size=self._format_file_size(file_size),
                file_path=file_path,
                user_id=user_id,
                metadata={
                    "mime_type": mime_type,
                    "file_size_bytes": file_size,
                    "upload_timestamp": datetime.now().isoformat()
                }
            )
            
            db.add(document)
            db.commit()
            db.refresh(document)
            
            return {
                "id": document.id,
                "name": document.name,
                "type": document.type,
                "size": document.size,
                "upload_date": document.upload_date.isoformat(),
                "message": "File đã được tải lên thành công"
            }
            
        except Exception as e:
            # Clean up file if database save fails
            if os.path.exists(file_path):
                os.remove(file_path)
            raise HTTPException(status_code=500, detail=f"Lỗi khi lưu file: {str(e)}")

    async def get_documents(
        self, 
        db: Session, 
        user_id: UUID, 
        page: int = 1, 
        limit: int = 10,
        search: Optional[str] = None,
        type_filter: Optional[str] = None
    ):
        """Lấy danh sách tài liệu của người dùng"""
        
        # Base query - documents owned by user or shared with user
        query = db.query(Document).filter(
            or_(
                Document.user_id == user_id,
                and_(
                    Document.shared == True,
                    Document.id.in_(
                        db.query(DocumentPermission.document_id).filter(
                            DocumentPermission.user_id == user_id
                        )
                    )
                )
            )
        )
        
        # Apply search filter
        if search:
            query = query.filter(
                or_(
                    Document.name.ilike(f"%{search}%"),
                    Document.extracted_text.ilike(f"%{search}%")
                )
            )
        
        # Apply type filter
        if type_filter:
            query = query.filter(Document.type == type_filter.upper())
        
        # Get total count
        total = query.count()
        
        # Apply pagination
        offset = (page - 1) * limit
        documents = query.offset(offset).limit(limit).all()
        
        return {
            "documents": [
                {
                    "id": doc.id,
                    "name": doc.name,
                    "type": doc.type,
                    "size": doc.size,
                    "upload_date": doc.upload_date.isoformat(),
                    "folder": doc.folder,
                    "shared": doc.shared,
                    "is_processed": doc.is_processed,
                    "is_owner": doc.user_id == user_id
                } for doc in documents
            ],
            "total": total,
            "page": page,
            "limit": limit,
            "pages": (total + limit - 1) // limit
        }

    async def delete_document(self, db: Session, document_id: str, user_id: UUID):
        """Xóa tài liệu"""
        
        document = db.query(Document).filter(
            and_(
                Document.id == document_id,
                Document.user_id == user_id  # Only owner can delete
            )
        ).first()
        
        if not document:
            raise HTTPException(status_code=404, detail="Không tìm thấy tài liệu hoặc bạn không có quyền xóa")
        
        try:
            # Delete physical file
            if os.path.exists(document.file_path):
                os.remove(document.file_path)
            
            # Delete from database
            db.delete(document)
            db.commit()
            
            return {"message": "Tài liệu đã được xóa thành công"}
            
        except Exception as e:
            db.rollback()
            raise HTTPException(status_code=500, detail=f"Lỗi khi xóa tài liệu: {str(e)}")

    async def share_document(self, db: Session, document_id: str, share_data: DocumentShare, user_id: UUID):
        """Chia sẻ tài liệu với người dùng khác"""
        
        # Check if document exists and user is owner
        document = db.query(Document).filter(
            and_(
                Document.id == document_id,
                Document.user_id == user_id
            )
        ).first()
        
        if not document:
            raise HTTPException(status_code=404, detail="Không tìm thấy tài liệu hoặc bạn không có quyền chia sẻ")
        
        # Find user to share with
        target_user = db.query(User).filter(User.email == share_data.user_email).first()
        if not target_user:
            raise HTTPException(status_code=404, detail="Không tìm thấy người dùng với email này")
        
        # Check if already shared
        existing_permission = db.query(DocumentPermission).filter(
            and_(
                DocumentPermission.document_id == document_id,
                DocumentPermission.user_id == target_user.id
            )
        ).first()
        
        if existing_permission:
            # Update existing permission
            existing_permission.permission = share_data.permission
            existing_permission.granted_at = datetime.now()
        else:
            # Create new permission
            permission = DocumentPermission(
                document_id=document.id,
                user_id=target_user.id,
                permission=share_data.permission,
                granted_by=user_id
            )
            db.add(permission)
        
        # Mark document as shared
        document.shared = True
        
        db.commit()
        
        return {
            "message": f"Tài liệu đã được chia sẻ với {share_data.user_email}",
            "permission": share_data.permission
        }

    async def get_document_content(self, db: Session, document_id: str, user_id: UUID):
        """Lấy nội dung tài liệu"""
        
        # Check if user has access to document
        document = db.query(Document).filter(
            and_(
                Document.id == document_id,
                or_(
                    Document.user_id == user_id,
                    and_(
                        Document.shared == True,
                        Document.id.in_(
                            db.query(DocumentPermission.document_id).filter(
                                DocumentPermission.user_id == user_id
                            )
                        )
                    )
                )
            )
        ).first()
        
        if not document:
            raise HTTPException(status_code=404, detail="Không tìm thấy tài liệu hoặc bạn không có quyền truy cập")
        
        try:
            if document.extracted_text:
                return {
                    "id": document.id,
                    "name": document.name,
                    "content": document.extracted_text,
                    "type": document.type
                }
            else:
                # If no extracted text, try to extract based on file type
                content = await self._extract_text_from_file(document.file_path, document.type)
                
                # Save extracted text for future use
                document.extracted_text = content
                document.is_processed = True
                db.commit()
                
                return {
                    "id": document.id,
                    "name": document.name,
                    "content": content,
                    "type": document.type
                }
                
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Lỗi khi đọc nội dung tài liệu: {str(e)}")

    async def _extract_text_from_file(self, file_path: str, file_type: str) -> str:
        """Trích xuất text từ file dựa vào loại file"""
        
        if file_type.upper() == 'PDF':
            return await self._extract_from_pdf(file_path)
        elif file_type.upper() in ['DOC', 'DOCX']:
            return await self._extract_from_docx(file_path)
        elif file_type.upper() == 'TXT':
            return await self._extract_from_txt(file_path)
        else:
            return "Không thể trích xuất text từ loại file này"

    async def _extract_from_pdf(self, file_path: str) -> str:
        """Trích xuất text từ PDF"""
        try:
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text.strip()
        except Exception as e:
            return f"Lỗi khi đọc PDF: {str(e)}"

    async def _extract_from_docx(self, file_path: str) -> str:
        """Trích xuất text từ DOCX"""
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            return f"Lỗi khi đọc DOCX: {str(e)}"

    async def _extract_from_txt(self, file_path: str) -> str:
        """Trích xuất text từ TXT"""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                return await f.read()
        except Exception as e:
            return f"Lỗi khi đọc TXT: {str(e)}"

    def _format_file_size(self, size_bytes: int) -> str:
        """Format file size to human readable string"""
        if size_bytes == 0:
            return "0 B"
        
        size_names = ["B", "KB", "MB", "GB"]
        i = 0
        size = float(size_bytes)
        
        while size >= 1024.0 and i < len(size_names) - 1:
            size /= 1024.0
            i += 1
        
        return f"{size:.1f} {size_names[i]}"