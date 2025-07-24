# app/services/search_service.py
import re
from fastapi import HTTPException
from sqlalchemy.orm import Session
from sqlalchemy import and_, or_, text
from typing import List, Dict, Any
from uuid import UUID
from datetime import datetime, timedelta

from app.models import Document, SearchHistory, DocumentPermission
from app.schemas import SearchRequest

class SearchService:
    def __init__(self):
        self.max_results = 100

    async def search(self, db: Session, search_data: SearchRequest, user_id: UUID):
        """Tìm kiếm tài liệu"""
        
        start_time = datetime.now()
        
        try:
            # Log search query
            await self._log_search(db, search_data.query, search_data.filters, user_id)
            
            # Build base query for user's accessible documents
            base_query = db.query(Document).filter(
                or_(
                    Document.user_id == user_id,
                    and_(
                        Document.shared == True,
                        Document.id.in_(
                            db.query(DocumentPermission.document_id).filter(
                                DocumentPermission.user_id == user_id
                            )
                        )
                    )
                )
            )
            
            # Apply search filters
            query = self._apply_search_filters(base_query, search_data)
            
            # Get total count
            total = query.count()
            
            # Apply pagination
            offset = (search_data.page - 1) * search_data.limit
            documents = query.offset(offset).limit(search_data.limit).all()
            
            # Process results
            results = []
            for doc in documents:
                result_item = await self._process_search_result(doc, search_data.query)
                results.append(result_item)
            
            # Calculate search time
            search_time = (datetime.now() - start_time).total_seconds()
            
            return {
                "results": results,
                "total": total,
                "page": search_data.page,
                "limit": search_data.limit,
                "query": search_data.query,
                "took": search_time
            }
            
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Lỗi khi tìm kiếm: {str(e)}")

    def _apply_search_filters(self, query, search_data: SearchRequest):
        """Áp dụng bộ lọc tìm kiếm"""
        
        # Text search in document name and content
        if search_data.query:
            search_terms = self._prepare_search_terms(search_data.query)
            text_conditions = []
            
            for term in search_terms:
                text_conditions.append(
                    or_(
                        Document.name.ilike(f"%{term}%"),
                        Document.extracted_text.ilike(f"%{term}%")
                    )
                )
            
            if text_conditions:
                query = query.filter(and_(*text_conditions))
        
        # Apply additional filters if provided
        if search_data.filters:
            query = self._apply_additional_filters(query, search_data.filters)
        
        # Order by relevance and date
        query = query.order_by(Document.upload_date.desc())
        
        return query

    def _apply_additional_filters(self, query, filters: Dict[str, Any]):
        """Áp dụng các bộ lọc bổ sung"""
        
        # Filter by document type
        if filters.get("type"):
            query = query.filter(Document.type == filters["type"].upper())
        
        # Filter by date range
        if filters.get("date_from"):
            try:
                date_from = datetime.fromisoformat(filters["date_from"])
                query = query.filter(Document.upload_date >= date_from)
            except:
                pass
        
        if filters.get("date_to"):
            try:
                date_to = datetime.fromisoformat(filters["date_to"])
                query = query.filter(Document.upload_date <= date_to)
            except:
                pass
        
        # Filter by folder
        if filters.get("folder"):
            query = query.filter(Document.folder == filters["folder"])
        
        # Filter by processing status
        if filters.get("processed") is not None:
            query = query.filter(Document.is_processed == filters["processed"])
        
        return query

    def _prepare_search_terms(self, query: str) -> List[str]:
        """Chuẩn bị các từ khóa tìm kiếm"""
        
        # Remove special characters and normalize
        clean_query = re.sub(r'[^\w\s]', ' ', query.lower())
        
        # Split into terms
        terms = [term.strip() for term in clean_query.split() if len(term.strip()) > 1]
        
        return terms[:10]  # Limit to 10 terms

    async def _process_search_result(self, document: Document, query: str) -> Dict[str, Any]:
        """Xử lý kết quả tìm kiếm để highlight và tạo excerpt"""
        
        highlights = []
        content_preview = ""
        
        if document.extracted_text:
            # Find highlights in content
            highlights = self._find_highlights(document.extracted_text, query)
            
            # Create content preview
            content_preview = self._create_content_preview(
                document.extracted_text, query
            )
        
        return {
            "id": document.id,
            "title": document.name,
            "content": content_preview,
            "author": "System",  # Will be replaced with actual user info
            "department": None,
            "date": document.upload_date,
            "type": document.type,
            "highlights": highlights,
            "score": 1.0  # Placeholder for relevance score
        }

    def _find_highlights(self, text: str, query: str) -> List[str]:
        """Tìm các đoạn text được highlight"""
        
        if not text or not query:
            return []
        
        highlights = []
        query_terms = self._prepare_search_terms(query)
        
        for term in query_terms:
            pattern = re.compile(re.escape(term), re.IGNORECASE)
            matches = pattern.finditer(text)
            
            for match in matches:
                start = max(0, match.start() - 50)
                end = min(len(text), match.end() + 50)
                highlight = text[start:end].strip()
                
                if highlight and highlight not in highlights:
                    highlights.append(highlight)
                
                if len(highlights) >= 3:  # Limit highlights
                    break
        
        return highlights

    def _create_content_preview(self, text: str, query: str) -> str:
        """Tạo preview nội dung với context"""
        
        if not text:
            return ""
        
        # If text is short, return as is
        if len(text) <= 300:
            return text
        
        # Try to find relevant excerpt based on query
        query_terms = self._prepare_search_terms(query)
        best_excerpt = ""
        best_score = 0
        
        # Split text into paragraphs
        paragraphs = text.split('\n')
        
        for paragraph in paragraphs:
            if len(paragraph.strip()) < 50:
                continue
            
            # Calculate relevance score
            score = 0
            for term in query_terms:
                score += paragraph.lower().count(term.lower())
            
            if score > best_score:
                best_score = score
                best_excerpt = paragraph.strip()
        
        # If no relevant excerpt found, use beginning
        if not best_excerpt:
            best_excerpt = text[:300].strip()
        
        # Truncate if too long
        if len(best_excerpt) > 300:
            best_excerpt = best_excerpt[:300] + "..."
        
        return best_excerpt

    async def get_suggestions(self, db: Session, query: str, user_id: UUID) -> List[str]:
        """Lấy gợi ý tìm kiếm"""
        
        if len(query) < 2:
            return []
        
        try:
            # Get recent searches
            recent_searches = db.query(SearchHistory.query).filter(
                and_(
                    SearchHistory.user_id == user_id,
                    SearchHistory.query.ilike(f"%{query}%")
                )
            ).distinct().limit(5).all()
            
            suggestions = [search.query for search in recent_searches]
            
            # Add document name suggestions
            document_suggestions = db.query(Document.name).filter(
                and_(
                    Document.user_id == user_id,
                    Document.name.ilike(f"%{query}%")
                )
            ).limit(5).all()
            
            suggestions.extend([doc.name for doc in document_suggestions])
            
            return list(set(suggestions))[:10]  # Remove duplicates and limit
            
        except Exception as e:
            return []

    async def _log_search(self, db: Session, query: str, filters: Dict[str, Any], user_id: UUID):
        """Log tìm kiếm"""
        
        try:
            search_log = SearchHistory(
                user_id=user_id,
                query=query,
                filters=filters,
                timestamp=datetime.now()
            )
            db.add(search_log)
            db.commit()
        except Exception as e:
            # Don't fail search if logging fails
            pass

# app/services/report_service.py
import os
import asyncio
from fastapi import HTTPException
from sqlalchemy.orm import Session
from sqlalchemy import and_
from typing import List, Dict, Any
from uuid import UUID
from datetime import datetime
from docx import Document as DocxDocument
from docx.shared import Inches
import json

from app.models import Report, Document, ChatMessage
from app.schemas import ReportCreate
from app.config import settings

class ReportService:
    def __init__(self):
        self.reports_dir = os.path.join(settings.UPLOAD_DIR, "reports")
        os.makedirs(self.reports_dir, exist_ok=True)

    async def generate_report(self, db: Session, report_config: ReportCreate, user_id: UUID):
        """Tạo báo cáo"""
        
        try:
            # Create report record
            report = Report(
                title=report_config.title,
                description=report_config.description,
                created_by=user_id,
                type=report_config.type,
                status="generating",
                config=report_config.config
            )
            
            db.add(report)
            db.commit()
            db.refresh(report)
            
            # Generate report asynchronously
            asyncio.create_task(self._generate_report_async(db, report.id, report_config, user_id))
            
            return {
                "id": report.id,
                "title": report.title,
                "status": "generating",
                "message": "Báo cáo đang được tạo..."
            }
            
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Lỗi khi tạo báo cáo: {str(e)}")

    async def _generate_report_async(self, db: Session, report_id: UUID, config: ReportCreate, user_id: UUID):
        """Tạo báo cáo bất đồng bộ"""
        
        try:
            report_content = ""
            file_path = None
            
            if config.type == "document_summary":
                report_content, file_path = await self._generate_document_summary(
                    db, config.config, user_id
                )
            elif config.type == "qa_analysis":
                report_content, file_path = await self._generate_qa_analysis(
                    db, config.config, user_id
                )
            elif config.type == "activity_report":
                report_content, file_path = await self._generate_activity_report(
                    db, config.config, user_id
                )
            else:
                raise ValueError(f"Unsupported report type: {config.type}")
            
            # Update report status
            report = db.query(Report).filter(Report.id == report_id).first()
            if report:
                report.content = report_content
                report.file_path = file_path
                report.status = "completed"
                db.commit()
                
        except Exception as e:
            # Update status to failed
            report = db.query(Report).filter(Report.id == report_id).first()
            if report:
                report.status = "failed"
                report.config = {**report.config, "error": str(e)}
                db.commit()

    async def _generate_document_summary(self, db: Session, config: Dict[str, Any], user_id: UUID) -> tuple[str, str]:
        """Tạo báo cáo tổng hợp tài liệu"""
        
        # Get documents based on config
        document_ids = config.get("document_ids", [])
        date_from = config.get("date_from")
        date_to = config.get("date_to")
        
        query = db.query(Document).filter(Document.user_id == user_id)
        
        if document_ids:
            query = query.filter(Document.id.in_(document_ids))
        
        if date_from:
            query = query.filter(Document.upload_date >= datetime.fromisoformat(date_from))
        
        if date_to:
            query = query.filter(Document.upload_date <= datetime.fromisoformat(date_to))
        
        documents = query.all()
        
        # Generate summary content
        content = f"# Báo cáo tổng hợp tài liệu\n\n"
        content += f"**Ngày tạo:** {datetime.now().strftime('%d/%m/%Y %H:%M')}\n\n"
        content += f"**Tổng số tài liệu:** {len(documents)}\n\n"
        
        # Document statistics
        type_stats = {}
        total_size = 0
        
        for doc in documents:
            doc_type = doc.type
            type_stats[doc_type] = type_stats.get(doc_type, 0) + 1
            
            # Try to get file size from metadata
            if doc.metadata and 'file_size_bytes' in doc.metadata:
                total_size += doc.metadata['file_size_bytes']
        
        content += "## Thống kê theo loại tài liệu\n\n"
        for doc_type, count in type_stats.items():
            content += f"- **{doc_type}:** {count} tài liệu\n"
        
        content += f"\n**Tổng dung lượng:** {self._format_file_size(total_size)}\n\n"
        
        # Document list
        content += "## Danh sách tài liệu\n\n"
        for doc in documents:
            content += f"### {doc.name}\n"
            content += f"- **Loại:** {doc.type}\n"
            content += f"- **Kích thước:** {doc.size}\n"
            content += f"- **Ngày tải lên:** {doc.upload_date.strftime('%d/%m/%Y %H:%M')}\n"
            
            if doc.extracted_text:
                preview = doc.extracted_text[:200] + "..." if len(doc.extracted_text) > 200 else doc.extracted_text
                content += f"- **Nội dung:** {preview}\n"
            
            content += "\n---\n\n"
        
        # Create DOCX file
        file_path = await self._create_docx_report(content, f"document_summary_{datetime.now().strftime('%Y%m%d_%H%M%S')}")
        
        return content, file_path

    async def _generate_qa_analysis(self, db: Session, config: Dict[str, Any], user_id: UUID) -> tuple[str, str]:
        """Tạo báo cáo phân tích Q&A"""
        
        # Get chat messages
        date_from = config.get("date_from")
        date_to = config.get("date_to")
        
        query = db.query(ChatMessage).join(ChatMessage.session).filter(
            ChatMessage.session.has(user_id=user_id)
        )
        
        if date_from:
            query = query.filter(ChatMessage.timestamp >= datetime.fromisoformat(date_from))
        
        if date_to:
            query = query.filter(ChatMessage.timestamp <= datetime.fromisoformat(date_to))
        
        messages = query.order_by(ChatMessage.timestamp.desc()).all()
        
        # Generate analysis content
        content = f"# Báo cáo phân tích hỏi đáp\n\n"
        content += f"**Ngày tạo:** {datetime.now().strftime('%d/%m/%Y %H:%M')}\n\n"
        
        questions = [msg for msg in messages if msg.type == "user"]
        answers = [msg for msg in messages if msg.type == "assistant"]
        
        content += f"**Tổng số câu hỏi:** {len(questions)}\n"
        content += f"**Tổng số câu trả lời:** {len(answers)}\n\n"
        
        # Recent questions
        content += "## Câu hỏi gần đây\n\n"
        for i, msg in enumerate(questions[:10]):
            content += f"### Câu hỏi {i+1}\n"
            content += f"**Thời gian:** {msg.timestamp.strftime('%d/%m/%Y %H:%M')}\n\n"
            content += f"**Nội dung:** {msg.content}\n\n"
            
            # Find corresponding answer
            answer = next((a for a in answers if a.timestamp > msg.timestamp), None)
            if answer:
                answer_preview = answer.content[:300] + "..." if len(answer.content) > 300 else answer.content
                content += f"**Trả lời:** {answer_preview}\n\n"
            
            content += "---\n\n"
        
        # Create DOCX file
        file_path = await self._create_docx_report(content, f"qa_analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}")
        
        return content, file_path

    async def _generate_activity_report(self, db: Session, config: Dict[str, Any], user_id: UUID) -> tuple[str, str]:
        """Tạo báo cáo hoạt động"""
        
        # Get user's documents and activities
        documents = db.query(Document).filter(Document.user_id == user_id).all()
        
        content = f"# Báo cáo hoạt động người dùng\n\n"
        content += f"**Ngày tạo:** {datetime.now().strftime('%d/%m/%Y %H:%M')}\n\n"
        
        # Document statistics
        content += "## Thống kê tài liệu\n\n"
        content += f"- **Tổng số tài liệu:** {len(documents)}\n"
        content += f"- **Tài liệu đã xử lý:** {sum(1 for doc in documents if doc.is_processed)}\n"
        content += f"- **Tài liệu được chia sẻ:** {sum(1 for doc in documents if doc.shared)}\n\n"
        
        # Recent uploads
        recent_docs = sorted(documents, key=lambda x: x.upload_date, reverse=True)[:10]
        content += "## Tài liệu tải lên gần đây\n\n"
        for doc in recent_docs:
            content += f"- **{doc.name}** ({doc.type}) - {doc.upload_date.strftime('%d/%m/%Y %H:%M')}\n"
        
        content += "\n"
        
        # Create DOCX file
        file_path = await self._create_docx_report(content, f"activity_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}")
        
        return content, file_path

    async def _create_docx_report(self, content: str, filename: str) -> str:
        """Tạo file DOCX từ nội dung markdown"""
        
        try:
            doc = DocxDocument()
            
            # Parse markdown content and add to document
            lines = content.split('\n')
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                if line.startswith('# '):
                    # Main heading
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    # Sub heading
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    # Sub sub heading
                    doc.add_heading(line[4:], level=3)
                elif line.startswith('**') and line.endswith('**'):
                    # Bold text
                    p = doc.add_paragraph()
                    p.add_run(line[2:-2]).bold = True
                elif line.startswith('- '):
                    # Bullet point
                    doc.add_paragraph(line[2:], style='List Bullet')
                elif line == '---':
                    # Separator
                    doc.add_paragraph('_' * 50)
                else:
                    # Regular paragraph
                    doc.add_paragraph(line)
            
            # Save file
            file_path = os.path.join(self.reports_dir, f"{filename}.docx")
            doc.save(file_path)
            
            return file_path
            
        except Exception as e:
            raise Exception(f"Lỗi khi tạo file DOCX: {str(e)}")

    async def get_reports(self, db: Session, user_id: UUID) -> List[dict]:
        """Lấy danh sách báo cáo"""
        
        reports = db.query(Report).filter(
            Report.created_by == user_id
        ).order_by(Report.created_date.desc()).all()
        
        return [
            {
                "id": report.id,
                "title": report.title,
                "description": report.description,
                "type": report.type,
                "status": report.status,
                "created_date": report.created_date.isoformat(),
                "has_file": bool(report.file_path)
            } for report in reports
        ]

    async def get_report(self, db: Session, report_id: str, user_id: UUID) -> dict:
        """Lấy chi tiết báo cáo"""
        
        report = db.query(Report).filter(
            and_(Report.id == report_id, Report.created_by == user_id)
        ).first()
        
        if not report:
            raise HTTPException(status_code=404, detail="Không tìm thấy báo cáo")
        
        return {
            "id": report.id,
            "title": report.title,
            "description": report.description,
            "type": report.type,
            "status": report.status,
            "content": report.content,
            "created_date": report.created_date.isoformat(),
            "config": report.config,
            "file_path": report.file_path
        }

    async def download_report(self, db: Session, report_id: str, user_id: UUID):
        """Tải báo cáo"""
        
        report = db.query(Report).filter(
            and_(Report.id == report_id, Report.created_by == user_id)
        ).first()
        
        if not report:
            raise HTTPException(status_code=404, detail="Không tìm thấy báo cáo")
        
        if not report.file_path or not os.path.exists(report.file_path):
            raise HTTPException(status_code=404, detail="File báo cáo không tồn tại")
        
        return {
            "file_path": report.file_path,
            "filename": f"{report.title}.docx",
            "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        }

    def _format_file_size(self, size_bytes: int) -> str:
        """Format file size to human readable string"""
        if size_bytes == 0:
            return "0 B"
        
        size_names = ["B", "KB", "MB", "GB"]
        i = 0
        size = float(size_bytes)
        
        while size >= 1024.0 and i < len(size_names) - 1:
            size /= 1024.0
            i += 1
        
        return f"{size:.1f} {size_names[i]}"